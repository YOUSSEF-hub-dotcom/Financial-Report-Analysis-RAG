"""
APIFileParser — routes uploaded files to the correct parser based on file extension.

HTML/TXT/SGML → existing html_table_parser.parse_sec_filing
PDF → LlamaParse (primary), pypdf (fallback)
DOCX → LlamaParse (primary), python-docx (fallback)
"""

import sys
import tempfile
from pathlib import Path
from typing import Any

from config.logging_config import get_logger
from config.settings import LLAMA_CLOUD_API_KEY

_src_root = Path(__file__).resolve().parent.parent.parent / "src"
for _p in (str(_src_root / "1_ingestion"), str(_src_root)):
    if _p not in sys.path:
        sys.path.insert(0, _p)

logger = get_logger("api.parsers")


class APIFileParser:
    """Routes uploaded files to the appropriate parser based on extension.

    Returns a dict with keys compatible with ``_ingest_document``:
        text_content  — raw text from the document
        tables        — list of Markdown-formatted table strings
        raw_html      — HTML content (populated only for HTML/SGML files)
    """

    SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({
        ".html", ".htm", ".txt", ".sgml", ".pdf", ".docx",
    })

    def __init__(self) -> None:
        self._llama_key: str = LLAMA_CLOUD_API_KEY or ""

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def parse_file(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        """Inspect file extension and dispatch to the appropriate parser."""
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{ext}'. "
                             f"Accepted: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}")

        logger.info("Parsing file: %s (type=%s, size=%d bytes)", filename, ext, len(file_bytes))

        if ext in (".html", ".htm", ".txt", ".sgml"):
            return self._parse_html_sgml(file_bytes, filename)
        if ext == ".pdf":
            return self._parse_pdf(file_bytes, filename)
        return self._parse_docx(file_bytes, filename)

    # ------------------------------------------------------------------
    # HTML / SGML / TXT — delegate to the existing SEC parser
    # ------------------------------------------------------------------

    def _parse_html_sgml(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        from html_table_parser import parse_sec_filing

        with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False, mode="wb") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            parsed = parse_sec_filing(tmp_path)
            return {
                "text_content": parsed["text_content"],
                "tables": parsed.get("tables", []),
                "raw_html": parsed.get("raw_html", ""),
            }
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    # ------------------------------------------------------------------
    # PDF — LlamaParse primary, pypdf fallback
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        if self._llama_key:
            try:
                return self._parse_with_llamaparse(file_bytes, filename)
            except Exception as exc:
                logger.warning("LlamaParse failed for %s, falling back to pypdf: %s", filename, exc)
        return self._parse_pdf_fallback(file_bytes)

    def _parse_pdf_fallback(self, file_bytes: bytes) -> dict[str, Any]:
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return {"text_content": raw_text, "tables": [], "raw_html": ""}

    # ------------------------------------------------------------------
    # DOCX — LlamaParse primary, python-docx fallback
    # ------------------------------------------------------------------

    def _parse_docx(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        if self._llama_key:
            try:
                return self._parse_with_llamaparse(file_bytes, filename)
            except Exception as exc:
                logger.warning("LlamaParse failed for %s, falling back to python-docx: %s", filename, exc)
        return self._parse_docx_fallback(file_bytes)

    def _parse_docx_fallback(self, file_bytes: bytes) -> dict[str, Any]:
        from io import BytesIO

        from docx import Document

        doc = Document(BytesIO(file_bytes))
        raw_text = "\n".join(p.text for p in doc.paragraphs)

        tables: list[str] = []
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                col_count = len(rows[0].split(" | "))
                sep = "| " + " | ".join(["---"] * col_count) + " |"
                tables.append("\n".join([sep, *rows]))

        return {"text_content": raw_text, "tables": tables, "raw_html": ""}

    # ------------------------------------------------------------------
    # LlamaParse wrapper (shared by PDF and DOCX)
    # ------------------------------------------------------------------

    def _parse_with_llamaparse(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        from llama_parse import LlamaParse

        parser = LlamaParse(api_key=self._llama_key, result_type="markdown", verbose=False)

        with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False, mode="wb") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            documents = parser.load_data([tmp_path])
            raw_text = "\n\n".join(doc.text for doc in documents) if documents else ""

            tables: list[str] = []
            non_table_lines: list[str] = []
            in_table = False
            table_buffer: list[str] = []
            for line in raw_text.split("\n"):
                stripped = line.strip()
                is_table = stripped.startswith("|") and stripped.endswith("|") and len(stripped) > 2
                if is_table:
                    in_table = True
                    table_buffer.append(line)
                else:
                    if in_table and table_buffer:
                        tables.append("\n".join(table_buffer))
                        table_buffer = []
                    in_table = False
                    non_table_lines.append(line)
            if table_buffer:
                tables.append("\n".join(table_buffer))

            return {"text_content": raw_text, "tables": tables, "raw_html": ""}
        finally:
            Path(tmp_path).unlink(missing_ok=True)
